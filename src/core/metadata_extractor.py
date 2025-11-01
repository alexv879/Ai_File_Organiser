"""
Advanced Metadata Extractor

Copyright (c) 2025 Alexandru Emanuel Vasile. All rights reserved.
Proprietary Software - 200-Key Limited Release License

This module extracts comprehensive metadata from files to enable better
organization decisions. Goes beyond basic file info to extract:
- EXIF data from images (GPS, camera, date/time)
- Document metadata (author, company, keywords, subject)
- Video metadata (duration, resolution, codec)
- Audio metadata (artist, album, genre, bitrate)
- PDF metadata (title, author, creation date, keywords)
- Office document properties

NOTICE: This software is proprietary and confidential.
See LICENSE.txt for full terms and conditions.

Author: Alexandru Emanuel Vasile
License: Proprietary (200-key limited release)
"""

import os
from pathlib import Path
from typing import Dict, Any, Optional, List
from datetime import datetime
import logging

logger = logging.getLogger(__name__)

# Optional dependencies - gracefully handle if not installed
try:
    from PIL import Image
    from PIL.ExifTags import TAGS, GPSTAGS
    HAS_PIL = True
except ImportError:
    HAS_PIL = False
    logger.warning("PIL/Pillow not installed - image metadata extraction disabled")

try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False
    logger.warning("PyPDF2 not installed - PDF metadata extraction disabled")

try:
    from mutagen import File as MutagenFile
    HAS_MUTAGEN = True
except ImportError:
    HAS_MUTAGEN = False
    logger.warning("Mutagen not installed - audio/video metadata extraction disabled")

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not installed - DOCX metadata extraction disabled")


class AdvancedMetadataExtractor:
    """
    Extracts comprehensive metadata from various file types.
    
    Supports:
    - Images: EXIF data, GPS location, camera info, date taken
    - PDFs: Author, title, subject, keywords, creation date
    - Office docs: Author, company, keywords, last modified by
    - Audio: Artist, album, genre, year, duration, bitrate
    - Video: Duration, resolution, codec, frame rate
    """
    
    def __init__(self):
        """Initialize metadata extractor"""
        self.supported_types = self._get_supported_types()
        logger.info(f"Initialized with support for: {', '.join(self.supported_types)}")
    
    def _get_supported_types(self) -> List[str]:
        """Get list of file types with metadata support"""
        types = ["basic"]  # Always available
        
        if HAS_PIL:
            types.append("images")
        if HAS_PYPDF2:
            types.append("pdf")
        if HAS_MUTAGEN:
            types.extend(["audio", "video"])
        if HAS_DOCX:
            types.append("docx")
        
        return types
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract comprehensive metadata from a file.
        
        Args:
            file_path: Path to file
            
        Returns:
            Dict with all extracted metadata
        """
        if not os.path.exists(file_path):
            return {"error": "File not found"}
        
        file_path = Path(file_path)
        extension = file_path.suffix.lower().lstrip('.')
        
        # Start with basic metadata (always available)
        metadata = self._extract_basic_metadata(file_path)
        
        # Add type-specific metadata
        if extension in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff']:
            metadata.update(self._extract_image_metadata(file_path))
        
        elif extension == 'pdf':
            metadata.update(self._extract_pdf_metadata(file_path))
        
        elif extension in ['mp3', 'flac', 'wav', 'aac', 'ogg', 'm4a', 'wma']:
            metadata.update(self._extract_audio_metadata(file_path))
        
        elif extension in ['mp4', 'avi', 'mkv', 'mov', 'wmv', 'flv', 'webm']:
            metadata.update(self._extract_video_metadata(file_path))
        
        elif extension == 'docx':
            metadata.update(self._extract_docx_metadata(file_path))
        
        # Add organization hints based on metadata
        metadata['organization_hints'] = self._generate_organization_hints(metadata)
        
        return metadata
    
    def _extract_basic_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract basic file system metadata (always available)"""
        stat = file_path.stat()
        
        return {
            'filename': file_path.name,
            'extension': file_path.suffix.lower().lstrip('.'),
            'size_bytes': stat.st_size,
            'size_mb': stat.st_size / (1024**2),
            'created_time': datetime.fromtimestamp(stat.st_ctime),
            'modified_time': datetime.fromtimestamp(stat.st_mtime),
            'accessed_time': datetime.fromtimestamp(stat.st_atime),
            'path': str(file_path),
            'parent_dir': file_path.parent.name
        }
    
    def _extract_image_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract EXIF metadata from images"""
        if not HAS_PIL:
            return {}
        
        metadata = {'metadata_type': 'image'}
        
        try:
            with Image.open(file_path) as img:
                metadata['dimensions'] = f"{img.width}x{img.height}"
                metadata['format'] = img.format
                metadata['mode'] = img.mode
                
                # Extract EXIF data
                exif_data = img._getexif()
                if exif_data:
                    exif = {}
                    
                    for tag_id, value in exif_data.items():
                        tag = TAGS.get(tag_id, tag_id)
                        
                        # Handle GPS data specially
                        if tag == 'GPSInfo':
                            gps_data = {}
                            for gps_tag in value:
                                gps_tag_name = GPSTAGS.get(gps_tag, gps_tag)
                                gps_data[gps_tag_name] = value[gps_tag]
                            metadata['gps'] = gps_data
                            
                            # Extract location hint
                            if 'GPSLatitude' in gps_data and 'GPSLongitude' in gps_data:
                                metadata['has_location'] = True
                        
                        # Store important EXIF tags
                        elif tag in ['DateTime', 'DateTimeOriginal', 'DateTimeDigitized']:
                            try:
                                metadata['date_taken'] = datetime.strptime(
                                    str(value), '%Y:%m:%d %H:%M:%S'
                                )
                            except:
                                metadata['date_taken_raw'] = str(value)
                        
                        elif tag in ['Make', 'Model']:  # Camera info
                            metadata[f'camera_{tag.lower()}'] = str(value)
                        
                        elif tag in ['Artist', 'Copyright']:
                            metadata[tag.lower()] = str(value)
                        
                        elif tag == 'ImageDescription':
                            metadata['description'] = str(value)
                        
                        elif tag == 'Software':
                            metadata['software'] = str(value)
                    
        except Exception as e:
            logger.error(f"Error extracting image metadata: {e}")
            metadata['extraction_error'] = str(e)
        
        return metadata
    
    def _extract_pdf_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from PDF files"""
        if not HAS_PYPDF2:
            return {}
        
        metadata = {'metadata_type': 'pdf'}
        
        try:
            with open(file_path, 'rb') as f:
                pdf = PyPDF2.PdfReader(f)
                
                metadata['pages'] = len(pdf.pages)
                
                # Extract document info
                info = pdf.metadata
                if info:
                    if info.title:
                        metadata['title'] = info.title
                    if info.author:
                        metadata['author'] = info.author
                    if info.subject:
                        metadata['subject'] = info.subject
                    if info.creator:
                        metadata['creator'] = info.creator
                    if info.producer:
                        metadata['producer'] = info.producer
                    
                    # Parse dates
                    if info.creation_date:
                        try:
                            metadata['creation_date'] = info.creation_date
                        except:
                            metadata['creation_date_raw'] = str(info.creation_date)
                    
                    if info.modification_date:
                        try:
                            metadata['modification_date'] = info.modification_date
                        except:
                            metadata['modification_date_raw'] = str(info.modification_date)
                
                # Try to extract text from first page for classification
                if len(pdf.pages) > 0:
                    try:
                        first_page_text = pdf.pages[0].extract_text()
                        metadata['first_page_preview'] = first_page_text[:500]
                        
                        # Detect document type from content
                        text_lower = first_page_text.lower()
                        if 'invoice' in text_lower:
                            metadata['detected_type'] = 'invoice'
                        elif 'receipt' in text_lower:
                            metadata['detected_type'] = 'receipt'
                        elif 'contract' in text_lower or 'agreement' in text_lower:
                            metadata['detected_type'] = 'contract'
                        elif 'report' in text_lower:
                            metadata['detected_type'] = 'report'
                    except:
                        pass
        
        except Exception as e:
            logger.error(f"Error extracting PDF metadata: {e}")
            metadata['extraction_error'] = str(e)
        
        return metadata
    
    def _extract_audio_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from audio files"""
        if not HAS_MUTAGEN:
            return {}
        
        metadata = {'metadata_type': 'audio'}
        
        try:
            audio = MutagenFile(str(file_path))
            if audio:
                # Duration
                if hasattr(audio.info, 'length'):
                    metadata['duration_seconds'] = int(audio.info.length)
                    metadata['duration_formatted'] = self._format_duration(audio.info.length)
                
                # Bitrate
                if hasattr(audio.info, 'bitrate'):
                    metadata['bitrate'] = audio.info.bitrate
                
                # Sample rate
                if hasattr(audio.info, 'sample_rate'):
                    metadata['sample_rate'] = audio.info.sample_rate
                
                # Tags
                if audio.tags:
                    # Common tags across formats
                    tag_mappings = {
                        'artist': ['artist', 'TPE1', '\xa9ART'],
                        'album': ['album', 'TALB', '\xa9alb'],
                        'title': ['title', 'TIT2', '\xa9nam'],
                        'genre': ['genre', 'TCON', '\xa9gen'],
                        'date': ['date', 'TDRC', '\xa9day'],
                        'comment': ['comment', 'COMM', '\xa9cmt']
                    }
                    
                    for std_name, possible_tags in tag_mappings.items():
                        for tag in possible_tags:
                            if tag in audio.tags:
                                value = audio.tags[tag]
                                if isinstance(value, list):
                                    value = value[0]
                                metadata[std_name] = str(value)
                                break
        
        except Exception as e:
            logger.error(f"Error extracting audio metadata: {e}")
            metadata['extraction_error'] = str(e)
        
        return metadata
    
    def _extract_video_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from video files"""
        if not HAS_MUTAGEN:
            return {}
        
        metadata = {'metadata_type': 'video'}
        
        try:
            video = MutagenFile(str(file_path))
            if video and hasattr(video, 'info'):
                # Duration
                if hasattr(video.info, 'length'):
                    metadata['duration_seconds'] = int(video.info.length)
                    metadata['duration_formatted'] = self._format_duration(video.info.length)
                
                # Resolution (if available)
                if hasattr(video.info, 'width') and hasattr(video.info, 'height'):
                    metadata['resolution'] = f"{video.info.width}x{video.info.height}"
                
                # Frame rate
                if hasattr(video.info, 'fps'):
                    metadata['fps'] = video.info.fps
                
                # Bitrate
                if hasattr(video.info, 'bitrate'):
                    metadata['bitrate'] = video.info.bitrate
        
        except Exception as e:
            logger.error(f"Error extracting video metadata: {e}")
            metadata['extraction_error'] = str(e)
        
        return metadata
    
    def _extract_docx_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from DOCX files"""
        if not HAS_DOCX:
            return {}
        
        metadata = {'metadata_type': 'docx'}
        
        try:
            doc = DocxDocument(str(file_path))
            core_props = doc.core_properties
            
            if core_props.author:
                metadata['author'] = core_props.author
            if core_props.category:
                metadata['category'] = core_props.category
            if core_props.comments:
                metadata['comments'] = core_props.comments
            if core_props.title:
                metadata['title'] = core_props.title
            if core_props.subject:
                metadata['subject'] = core_props.subject
            if core_props.keywords:
                metadata['keywords'] = core_props.keywords
            if core_props.last_modified_by:
                metadata['last_modified_by'] = core_props.last_modified_by
            if core_props.created:
                metadata['created'] = core_props.created
            if core_props.modified:
                metadata['modified'] = core_props.modified
            
            # Extract first paragraph for classification
            if doc.paragraphs:
                text = ' '.join([p.text for p in doc.paragraphs[:3]])
                metadata['content_preview'] = text[:500]
        
        except Exception as e:
            logger.error(f"Error extracting DOCX metadata: {e}")
            metadata['extraction_error'] = str(e)
        
        return metadata
    
    def _format_duration(self, seconds: float) -> str:
        """Format duration in HH:MM:SS"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        
        if hours > 0:
            return f"{hours:02d}:{minutes:02d}:{secs:02d}"
        else:
            return f"{minutes:02d}:{secs:02d}"
    
    def _generate_organization_hints(self, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """
        Generate organization hints based on extracted metadata.
        
        Returns suggestions for better file organization.
        """
        hints = {}
        
        # Image-specific hints
        if metadata.get('metadata_type') == 'image':
            if metadata.get('has_location'):
                hints['location_based'] = True
                hints['suggestion'] = "Consider organizing by location using GPS data"
            
            if 'date_taken' in metadata:
                hints['date_based'] = True
                date = metadata['date_taken']
                hints['suggested_date_folder'] = f"{date.year}/{date.strftime('%B')}"
            
            if 'camera_make' in metadata:
                hints['camera_based'] = True
                hints['suggested_camera_folder'] = f"Camera-{metadata['camera_make']}"
        
        # PDF-specific hints
        elif metadata.get('metadata_type') == 'pdf':
            if metadata.get('author'):
                hints['author_based'] = True
                hints['suggested_author_folder'] = f"Author-{metadata['author']}"
            
            if metadata.get('detected_type'):
                hints['type_detected'] = metadata['detected_type']
                hints['confidence'] = 'high'
            
            if metadata.get('subject'):
                hints['subject_based'] = True
        
        # Audio-specific hints
        elif metadata.get('metadata_type') == 'audio':
            if metadata.get('artist') and metadata.get('album'):
                hints['music_organization'] = True
                hints['suggested_path'] = f"Music/{metadata['artist']}/{metadata['album']}"
            
            if metadata.get('genre'):
                hints['genre_based'] = True
                hints['suggested_genre_folder'] = f"Genre-{metadata['genre']}"
            
            if metadata.get('date'):
                hints['year_based'] = True
        
        # Video-specific hints
        elif metadata.get('metadata_type') == 'video':
            if metadata.get('resolution'):
                hints['resolution_based'] = True
                # Suggest 4K, HD, SD folders
                width = int(metadata['resolution'].split('x')[0])
                if width >= 3840:
                    hints['quality_folder'] = '4K-Ultra-HD'
                elif width >= 1920:
                    hints['quality_folder'] = 'Full-HD'
                elif width >= 1280:
                    hints['quality_folder'] = 'HD'
                else:
                    hints['quality_folder'] = 'Standard'
            
            if metadata.get('duration_seconds'):
                duration = metadata['duration_seconds']
                if duration < 60:  # Short clips
                    hints['duration_category'] = 'Short-Clips'
                elif duration < 600:  # Less than 10 minutes
                    hints['duration_category'] = 'Medium-Videos'
                else:
                    hints['duration_category'] = 'Long-Videos'
        
        # Document-specific hints
        elif metadata.get('metadata_type') == 'docx':
            if metadata.get('author'):
                hints['author_based'] = True
            
            if metadata.get('category'):
                hints['category_based'] = True
                hints['suggested_category'] = metadata['category']
            
            if metadata.get('keywords'):
                hints['keywords'] = metadata['keywords']
        
        # Date-based hints (for all types)
        if 'modified_time' in metadata:
            mod_time = metadata['modified_time']
            age_days = (datetime.now() - mod_time).days
            
            if age_days > 365:
                hints['is_old'] = True
                hints['archive_candidate'] = True
                hints['suggested_archive_year'] = mod_time.year
            elif age_days > 180:
                hints['is_mature'] = True
        
        return hints


def create_metadata_extractor() -> AdvancedMetadataExtractor:
    """Create a metadata extractor instance"""
    return AdvancedMetadataExtractor()


if __name__ == "__main__":
    # Test metadata extractor
    print("Testing Advanced Metadata Extractor...\n")
    
    extractor = AdvancedMetadataExtractor()
    
    print(f"Supported types: {', '.join(extractor.supported_types)}\n")
    
    # Test with a sample file (if exists)
    test_file = "D:\\AIFILEORGANISER\\README.md"
    if os.path.exists(test_file):
        print(f"Testing with: {test_file}\n")
        metadata = extractor.extract_metadata(test_file)
        
        print("Extracted Metadata:")
        print("-" * 50)
        for key, value in metadata.items():
            if key != 'organization_hints':
                print(f"{key}: {value}")
        
        print("\nOrganization Hints:")
        print("-" * 50)
        hints = metadata.get('organization_hints', {})
        for key, value in hints.items():
            print(f"{key}: {value}")
    else:
        print("No test file found")
